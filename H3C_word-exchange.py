import os
import re
from tkinter import *
from tkinter import ttk, filedialog, messagebox
import threading
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
from openpyxl.utils import get_column_letter


class WordToExcelConverter:
    def __init__(self, root):
        self.root = root
        self.root.title("Word漏洞报告转换器")
        self.root.geometry("700x600")

        # 存储选择的文件路径
        self.word_files = []
        self.output_path = ""

        self.setup_ui()

    def setup_ui(self):
        # 标题
        title_label = Label(self.root, text="Word漏洞报告转换器", font=("Arial", 16, "bold"))
        title_label.pack(pady=10)

        # 选择Word文件区域
        file_frame = Frame(self.root)
        file_frame.pack(fill="x", padx=20, pady=10)

        Label(file_frame, text="选择Word文件:", font=("Arial", 10)).pack(anchor="w")

        # 文件列表显示
        self.file_listbox = Listbox(file_frame, height=5, selectmode=EXTENDED)
        self.file_listbox.pack(fill="x", pady=5)

        # 文件操作按钮
        file_button_frame = Frame(file_frame)
        file_button_frame.pack(fill="x")

        Button(file_button_frame, text="添加文件", command=self.add_files, width=12).pack(side="left", padx=2)
        Button(file_button_frame, text="添加文件夹", command=self.add_folder, width=12).pack(side="left", padx=2)
        Button(file_button_frame, text="清空列表", command=self.clear_files, width=12).pack(side="left", padx=2)

        # 选择输出路径区域
        output_frame = Frame(self.root)
        output_frame.pack(fill="x", padx=20, pady=10)

        Label(output_frame, text="输出Excel文件:", font=("Arial", 10)).pack(anchor="w")

        output_path_frame = Frame(output_frame)
        output_path_frame.pack(fill="x", pady=5)

        self.output_path_var = StringVar()
        Entry(output_path_frame, textvariable=self.output_path_var, state='readonly').pack(side="left", fill="x",
                                                                                           expand=True)
        Button(output_path_frame, text="浏览", command=self.select_output_file, width=10).pack(side="left", padx=5)

        # 转换按钮
        self.convert_button = Button(self.root, text="开始转换", command=self.start_conversion,
                                     bg="#4CAF50", fg="white", font=("Arial", 12), padx=20, pady=10)
        self.convert_button.pack(pady=15)

        # 进度条
        self.progress = ttk.Progressbar(self.root, length=500, mode='indeterminate')

        # 状态显示
        self.status_label = Label(self.root, text="", fg="blue")
        self.status_label.pack(pady=5)

        # 标签页
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill="both", expand=True, padx=20, pady=10)

        # 转换日志标签页
        log_frame = Frame(self.notebook)
        self.notebook.add(log_frame, text="转换日志")

        self.log_text = Text(log_frame, height=10, width=80)
        scrollbar = Scrollbar(log_frame)
        scrollbar.pack(side=RIGHT, fill=Y)
        self.log_text.pack(side=LEFT, fill="both", expand=True)
        scrollbar.config(command=self.log_text.yview)
        self.log_text.config(yscrollcommand=scrollbar.set)

    def add_files(self):
        files = filedialog.askopenfilenames(
            title="选择Word文档",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if files:
            for file in files:
                if file not in self.word_files:
                    self.word_files.append(file)
                    self.file_listbox.insert(END, os.path.basename(file))
            self.log_message(f"已添加 {len(files)} 个文件")

    def add_folder(self):
        folder = filedialog.askdirectory(title="选择文件夹")
        if folder:
            added_count = 0
            for file in os.listdir(folder):
                if file.endswith('.docx'):
                    file_path = os.path.join(folder, file)
                    if file_path not in self.word_files:
                        self.word_files.append(file_path)
                        self.file_listbox.insert(END, file)
                        added_count += 1
            if added_count > 0:
                self.log_message(f"从文件夹添加了 {added_count} 个文件")

    def clear_files(self):
        self.word_files.clear()
        self.file_listbox.delete(0, END)
        self.log_message("文件列表已清空")

    def select_output_file(self):
        file_path = filedialog.asksaveasfilename(
            title="保存Excel文件",
            defaultextension=".xlsx",
            filetypes=[("Excel文件", "*.xlsx"), ("所有文件", "*.*")]
        )
        if file_path:
            self.output_path = file_path
            self.output_path_var.set(file_path)

    def log_message(self, message):
        self.status_label.config(text=message)
        self.log_text.insert(END, f"{message}\n")
        self.log_text.see(END)
        self.root.update()

    def start_conversion(self):
        if not self.word_files:
            messagebox.showwarning("警告", "请先选择Word文件！")
            return

        if not self.output_path:
            messagebox.showwarning("警告", "请选择输出文件路径！")
            return

        # 禁用转换按钮，开始进度条
        self.convert_button.config(state=DISABLED)
        self.progress.pack(pady=5)
        self.progress.start()

        # 在新线程中执行转换
        thread = threading.Thread(target=self.convert_files)
        thread.start()

    def convert_files(self):
        try:
            total_files = len(self.word_files)
            self.log_message(f"开始处理 {total_files} 个Word文件...")

            # 存储所有IP的漏洞信息
            all_vulnerabilities = {}
            ip_with_high_or_medium_risk = []
            ip_high_risk_counts = {}

            for i, word_file in enumerate(self.word_files, 1):
                try:
                    self.log_message(f"正在处理文件 {i}/{total_files}: {os.path.basename(word_file)}")

                    # 提取IP地址
                    filename = os.path.basename(word_file)
                    ip_match = re.match(r'^([\d\.]+)-', filename)
                    if ip_match:
                        ip_address = ip_match.group(1)
                    else:
                        ip_address = f"未知IP_{i}"

                    # 处理Word文档
                    vulnerabilities = self.extract_vulnerabilities_from_tables(word_file)

                    if vulnerabilities:
                        # 过滤漏洞，只保留包含"高风险"或"中风险"的漏洞
                        filtered_vulnerabilities = []
                        high_risk_count = 0
                        medium_risk_count = 0

                        for vuln in vulnerabilities:
                            risk_level = vuln.get('风险等级', '')
                            # 检查是否包含"高风险"或"中风险"
                            if "高风险" in risk_level or "中风险" in risk_level:
                                filtered_vulnerabilities.append(vuln)
                                # 统计高风险数量
                                if "高风险" in risk_level:
                                    high_risk_count += 1
                                elif "中风险" in risk_level:
                                    medium_risk_count += 1

                        if filtered_vulnerabilities:
                            all_vulnerabilities[ip_address] = filtered_vulnerabilities
                            ip_with_high_or_medium_risk.append(ip_address)
                            ip_high_risk_counts[ip_address] = high_risk_count
                            self.log_message(
                                f"✓ {ip_address} 找到 {len(filtered_vulnerabilities)} 个漏洞 (高风险: {high_risk_count}, 中风险: {medium_risk_count})")
                        else:
                            self.log_message(f"⚠ {ip_address} 没有高风险或中风险漏洞，跳过")
                    else:
                        self.log_message(f"⚠ {ip_address} 未找到漏洞数据")

                except Exception as e:
                    self.log_message(f"✗ 处理失败 {filename}: {str(e)}")

            # 创建Excel文件
            if ip_with_high_or_medium_risk:
                # 按高风险数量降序排序IP列表
                sorted_ips = sorted(ip_with_high_or_medium_risk,
                                    key=lambda ip: ip_high_risk_counts.get(ip, 0),
                                    reverse=True)

                self.create_excel_with_summary(self.output_path, all_vulnerabilities, sorted_ips, ip_high_risk_counts)
                self.log_message(f"✓ 成功创建Excel文件: {os.path.basename(self.output_path)}")

                # 显示排序后的IP列表
                sorted_info = []
                for ip in sorted_ips:
                    count = ip_high_risk_counts.get(ip, 0)
                    sorted_info.append(f"{ip}(高风险:{count})")
                self.log_message(f"✓ IP排序结果: {', '.join(sorted_info)}")
            else:
                self.log_message("⚠ 没有找到任何包含高风险或中风险漏洞的IP，不生成Excel文件")
                messagebox.showinfo("完成", "没有找到任何包含高风险或中风险漏洞的IP，未生成Excel文件")

            # 完成处理
            self.progress.stop()
            self.progress.pack_forget()
            self.convert_button.config(state=NORMAL)

            if ip_with_high_or_medium_risk:
                messagebox.showinfo("完成",
                                    f"转换完成！\n成功处理了 {len(ip_with_high_or_medium_risk)}/{total_files} 个IP\n已保存到: {os.path.basename(self.output_path)}")

        except Exception as e:
            self.log_message(f"✗ 转换过程出错: {str(e)}")
            self.progress.stop()
            self.progress.pack_forget()
            self.convert_button.config(state=NORMAL)
            messagebox.showerror("错误", f"转换过程出错: {str(e)}")

    def extract_vulnerabilities_from_tables(self, word_file):
        """从Word文档的表格中提取漏洞信息"""
        doc = Document(word_file)
        vulnerabilities = []

        # 查找所有"主机漏洞详情"的位置
        host_vuln_indices = []
        for i, para in enumerate(doc.paragraphs):
            if "主机漏洞详情" in para.text:
                host_vuln_indices.append(i)

        if len(host_vuln_indices) < 2:
            return vulnerabilities

        # 使用第二个"主机漏洞详情"的位置
        second_host_vuln_index = host_vuln_indices[1]

        # 查找"Web漏洞信息"的位置
        web_vuln_index = -1
        for i, para in enumerate(doc.paragraphs):
            if "Web漏洞信息" in para.text:
                web_vuln_index = i
                break

        if web_vuln_index == -1:
            return vulnerabilities

        # 获取文档中的所有表格
        all_tables = doc.tables

        # 遍历每个表格
        for table_idx, table in enumerate(all_tables):
            # 提取当前表格中的漏洞信息
            table_vulns = self.extract_vuln_from_table(table, table_idx)

            # 将提取到的漏洞添加到总列表
            if table_vulns:
                vulnerabilities.extend(table_vulns)

        return vulnerabilities

    def extract_vuln_from_table(self, table, table_idx):
        """从单个表格中提取漏洞信息"""
        vulnerabilities = []
        current_vuln = {}
        target_fields = ["漏洞名称", "风险等级", "解决办法"]

        # 标记是否已经开始提取当前漏洞
        extracting = False

        # 遍历表格的每一行
        for row_idx, row in enumerate(table.rows):
            # 检查每一列，查找目标字段
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()

                # 检查是否为目标字段
                for field in target_fields:
                    if field in cell_text:
                        # 如果找到"漏洞名称"，开始一个新的漏洞记录
                        if field == "漏洞名称":
                            # 如果已经有正在提取的漏洞，保存它
                            if current_vuln and extracting:
                                if all(f in current_vuln for f in ["漏洞名称", "风险等级", "解决办法"]):
                                    vulnerabilities.append(current_vuln.copy())

                            # 开始新的漏洞记录
                            current_vuln = {}
                            extracting = True

                        # 提取右侧单元格的内容
                        if col_idx + 1 < len(row.cells):
                            right_cell = row.cells[col_idx + 1]
                            content = right_cell.text.strip()
                            current_vuln[field] = content

                        # 如果找到"解决办法"，保存当前漏洞并停止遍历当前表格
                        if field == "解决办法":
                            if current_vuln and all(f in current_vuln for f in ["漏洞名称", "风险等级", "解决办法"]):
                                vulnerabilities.append(current_vuln.copy())

                            # 停止遍历当前表格
                            return vulnerabilities

        # 如果表格遍历结束，检查最后一个漏洞
        if current_vuln and all(f in current_vuln for f in ["漏洞名称", "风险等级", "解决办法"]):
            vulnerabilities.append(current_vuln.copy())

        return vulnerabilities

    def create_excel_with_summary(self, excel_path, all_vulnerabilities, sorted_ips, ip_high_risk_counts):
        """创建Excel文件，包含汇总页和各IP的分页"""
        wb = Workbook()

        # 删除默认的Sheet
        if "Sheet" in wb.sheetnames:
            wb.remove(wb["Sheet"])

        # 创建汇总页
        summary_ws = wb.create_sheet(title="漏洞汇总")

        # 设置列标题
        headers = ["IP地址", "漏洞名称", "风险等级", "解决办法"]
        for col, header in enumerate(headers, 1):
            cell = summary_ws.cell(row=1, column=col, value=header)
            cell.alignment = Alignment(horizontal='center', vertical='center')

            # 设置标题字体为粗体
            original_font = cell.font
            cell.font = Font(
                name=original_font.name,
                size=original_font.size,
                bold=True,
                italic=original_font.italic,
                vertAlign=original_font.vertAlign,
                underline=original_font.underline,
                strike=original_font.strike,
                color=original_font.color
            )

        current_row = 2

        # 预定义的颜色列表（6个颜色循环使用）
        color_list = [
            "8EA9DB",  # 蓝色，个性色1，淡色60%
            "F4B084",  # 红色，个性色2，淡色60%
            "C5E0B4",  # 橄榄色，个性色3，淡色60%
            "B4C6E7",  # 紫色，个性色4，淡色60%
            "A8D8E7",  # 水绿色，个性色5，淡色60%
            "FFD966",  # 橙色，个性色6，淡色60%
        ]

        # 为每个IP分配颜色
        ip_colors = {}
        for idx, ip_address in enumerate(sorted_ips):
            color_index = idx % len(color_list)
            ip_colors[ip_address] = PatternFill(fill_type="solid", start_color=color_list[color_index],
                                                end_color=color_list[color_index])

        # 定义边框样式（内外边框）
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        # 写入汇总数据，按照排序后的IP顺序
        for ip_address in sorted_ips:
            vulnerabilities = all_vulnerabilities.get(ip_address, [])

            for vuln in vulnerabilities:
                risk_level = vuln.get('风险等级', '')
                # 转换风险等级：包含"中风险"的显示为"中风险"，包含"高风险"的显示为"高风险"
                if "中风险" in risk_level:
                    display_risk = "中风险"
                elif "高风险" in risk_level:
                    display_risk = "高风险"
                else:
                    display_risk = risk_level

                summary_ws.cell(row=current_row, column=1, value=ip_address)
                summary_ws.cell(row=current_row, column=2, value=vuln.get('漏洞名称', ''))
                summary_ws.cell(row=current_row, column=3, value=display_risk)
                summary_ws.cell(row=current_row, column=4, value=vuln.get('解决办法', ''))

                # 为整行应用颜色
                for col in range(1, 5):
                    cell = summary_ws.cell(row=current_row, column=col)
                    cell.fill = ip_colors[ip_address]
                    # 为数据部分添加边框
                    cell.border = thin_border

                current_row += 1

        # 为标题行添加边框
        for col in range(1, 5):
            summary_ws.cell(row=1, column=col).border = thin_border

        # 设置汇总页的行高和列宽
        for row in summary_ws.iter_rows(min_row=1, max_row=summary_ws.max_row):
            summary_ws.row_dimensions[row[0].row].height = 20

        for col in range(1, 5):
            summary_ws.column_dimensions[get_column_letter(col)].width = 30

        # 设置汇总页单元格对齐方式
        for row in summary_ws.iter_rows(min_row=2, max_row=summary_ws.max_row):
            for cell in row:
                cell.alignment = Alignment(vertical='center', wrap_text=True)

        # 创建各IP的分页，按照排序后的IP顺序
        for ip_address in sorted_ips:
            vulnerabilities = all_vulnerabilities.get(ip_address, [])

            # 创建IP的sheet页，只使用IP地址命名
            sheet_title = str(ip_address)
            # Excel sheet名称不能超过31个字符
            if len(sheet_title) > 31:
                sheet_title = sheet_title[:31]

            ws = wb.create_sheet(title=sheet_title)

            # 设置列标题
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.alignment = Alignment(horizontal='center', vertical='center')

                # 设置标题字体为粗体
                original_font = cell.font
                cell.font = Font(
                    name=original_font.name,
                    size=original_font.size,
                    bold=True,
                    italic=original_font.italic,
                    vertAlign=original_font.vertAlign,
                    underline=original_font.underline,
                    strike=original_font.strike,
                    color=original_font.color
                )
                # 为标题行添加边框
                cell.border = thin_border

            # 写入数据
            for row_idx, vuln in enumerate(vulnerabilities, 2):
                risk_level = vuln.get('风险等级', '')
                # 转换风险等级：包含"中风险"的显示为"中风险"，包含"高风险"的显示为"高风险"
                if "中风险" in risk_level:
                    display_risk = "中风险"
                elif "高风险" in risk_level:
                    display_risk = "高风险"
                else:
                    display_risk = risk_level

                ws.cell(row=row_idx, column=1, value=ip_address)
                ws.cell(row=row_idx, column=2, value=vuln.get('漏洞名称', ''))
                ws.cell(row=row_idx, column=3, value=display_risk)
                ws.cell(row=row_idx, column=4, value=vuln.get('解决办法', ''))

                # 为数据部分添加边框
                for col in range(1, 5):
                    ws.cell(row=row_idx, column=col).border = thin_border

            # 设置行高和列宽
            for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
                ws.row_dimensions[row[0].row].height = 20

            for col in range(1, 5):
                ws.column_dimensions[get_column_letter(col)].width = 30

            # 设置单元格对齐方式
            for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
                for cell in row:
                    cell.alignment = Alignment(vertical='center', wrap_text=True)

        # 保存文件
        wb.save(excel_path)


def main():
    root = Tk()
    app = WordToExcelConverter(root)
    root.mainloop()


if __name__ == "__main__":
    main()